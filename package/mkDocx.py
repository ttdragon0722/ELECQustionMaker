from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from os.path import abspath
from json import dump

def set_col_widths(table):
    widths = (Inches(2), Inches(8))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

# 製作輸出word檔案

def QuestionDocx(path,data:list,bank:list):
    document = Document()

    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    # 設置表格的樣式
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'answer'
    hdr_cells[1].text = 'question'

    # 設置表格的邊框
    table.border_color = RGBColor(0, 0, 0)
    table.border_width = 0.5

    for index,QuestionNum in enumerate(data,start=1):
        QuestionText = f"""第 {index} 題:
{bank[QuestionNum][1]}
A.{bank[QuestionNum][5]}
B.{bank[QuestionNum][6]}
C.{bank[QuestionNum][7]}
D.{bank[QuestionNum][8]}

{"" if bank[QuestionNum][4]=='' else bank[QuestionNum][4]}

解析:
{bank[QuestionNum][3]}"""
        row_cells = table.add_row().cells
        row_cells[0].width = Inches(0.5)


        # answers = [bank[QuestionNum][5],bank[QuestionNum][6],bank[QuestionNum][7],bank[QuestionNum][8]]
        # match answers.index(bank[QuestionNum][2]): 
        #     case 0:
        #         answer = "(A)"
        #     case 1:
        #         answer = "(B)"
        #     case 2:
        #         answer = "(C)"
        #     case 3:
        #         answer = "(D)"

        row_cells[0].text = str(bank[QuestionNum][2])
        row_cells[1].text = QuestionText

    set_col_widths(table)
    document.save(path+"\question.docx")    

def AnswerDocx(path,data:list,bank:list):
    document = Document()

    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    # 設置表格的樣式
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'question'

    # 設置表格的邊框
    table.border_color = RGBColor(0, 0, 0)
    table.border_width = 0.5

    for index,QuestionNum in enumerate(data,start=1):
        QuestionText = f"""第 {index} 題:
{bank[QuestionNum][1]}
A.{bank[QuestionNum][5]}
B.{bank[QuestionNum][6]}
C.{bank[QuestionNum][7]}
D.{bank[QuestionNum][8]}

{"" if bank[QuestionNum][4]=='' else bank[QuestionNum][4]}
"""
        row_cells = table.add_row().cells
        row_cells[0].width = Inches(0.5)


        # answers = [bank[QuestionNum][5],bank[QuestionNum][6],bank[QuestionNum][7],bank[QuestionNum][8]]
        # match answers.index(bank[QuestionNum][2]): 
        #     case 0:
        #         answer = "(A)"
        #     case 1:
        #         answer = "(B)"
        #     case 2:
        #         answer = "(C)"
        #     case 3:
        #         answer = "(D)"
        row_cells[0].text = QuestionText
    document.save(path+'\\answer.docx')

def MkJson(path,data,bank):
    result = {}
    for idx,d in enumerate(data,start=1):
        if bank[d][4]:
            dic = {
                'number':bank[d][0],
                'question':bank[d][1],
                'img':abspath(bank[d][4]),
                'options':[
                    bank[d][5],
                    bank[d][6],
                    bank[d][7],
                    bank[d][8]
                ],
                "answer":bank[d][2],
                'analyze':bank[d][3]
            }
        else:
            dic = {
                'number':bank[d][0],
                'question':bank[d][1],
                'img':"",
                'options':[
                    bank[d][5],
                    bank[d][6],
                    bank[d][7],
                    bank[d][8]
                ],
                "answer":bank[d][2],
                'analyze':bank[d][3]
            }
        result.update({idx:dic})

    with open(path+'\output.json',mode='w',encoding='utf-8') as f:
        dump(result,f,ensure_ascii=False)
